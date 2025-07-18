import tkinter as tk
from tkinter import ttk, messagebox
from pai_clinical_text import PAI_PARAGRAPHS
from generate_report_together import record_audio, transcribe_audio, clean_with_together
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import queue
import threading
import sounddevice as sd
import numpy as np
import scipy.io.wavfile as wav
import wave
import re

ALL_TESTS = [
    "*Wechsler Abbreviated Scale of Intelligence-II (WASI-II)",
    "*Portion of Wechsler Adult Intelligence Scale - Revised (WAIS-R)",
    "Portion of Wechsler Intelligence Scale for Children (WISC)",
    "*Trail Making Test (Part B)",
    "*Letter-Number Sequencing",
    "*BAARS-IV",
    "*Personality Assessment Inventory (PAI)",
    "Personality Assessment Inventory-Adolescent (PAI-A)",
    "Millon Clinical Multiaxial Inventory-III (MCMI-III)",
    "Millon Adolescent Clinical Inventory (MACI)",
    "Millon Pre-Adolescent Clinical Inventory (M-PACI)",
    "Minnesota Multiphasic Personality Inventory-2-RF (MMPI-2-RF)",
    "Minnesota Multiphasic Personality Inventory-Adolescent-RF (MMPI-A-RF)",
    "Rorschach Inkblot Test",
    "Thematic Apperception Test (TAT)",
    "Figure Drawing"
]

DEFAULT_SELECTED = [i for i, t in enumerate(ALL_TESTS) if t.startswith("*")]

def superscript_ordinals_in_doc(doc):
    import re
    ordinal_pattern = re.compile(r'(\d+)(st|nd|rd|th)')

    for para in doc.paragraphs:
        # We’ll rebuild runs carefully to preserve font style
        new_runs = []

        for run in para.runs:
            text = run.text
            matches = list(ordinal_pattern.finditer(text))
            if not matches:
                new_runs.append(run)
                continue

            # Clear original run text — we’ll rebuild runs below
            run.text = ''

            last_end = 0
            for m in matches:
                # Text before match
                if m.start() > last_end:
                    pre_text = text[last_end:m.start()]
                    r = para.add_run(pre_text)
                    copy_font(run, r)

                # Number part (normal)
                number = m.group(1)
                r_num = para.add_run(number)
                copy_font(run, r_num)

                # Ordinal suffix part (superscript)
                suffix = m.group(2)
                r_sup = para.add_run(suffix)
                copy_font(run, r_sup)
                r_sup.font.superscript = True

                last_end = m.end()

            # Text after last match
            if last_end < len(text):
                post_text = text[last_end:]
                r = para.add_run(post_text)
                copy_font(run, r)

        # Replace runs with new_runs if you want,
        # but python-docx does not allow run deletion easily,
        # so this approach is a little hacky but usually works.

def copy_font(source_run, target_run):
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.bold = source_run.font.bold
    target_run.font.italic = source_run.font.italic
    target_run.font.underline = source_run.font.underline
    # Add other font properties if needed

def ordinal(n):
    n = int(n)
    if 10 <= n % 100 <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
    return f"{n}{suffix}"

def clean_ordered_by(name):
    # Split into parts (e.g., "dr. jane smith md")
    parts = name.strip().split()
    corrected_parts = []

    for part in parts:
        # Keep known credentials all caps
        if part.upper() in {"MD", "NP", "APRN", "PA", "DO", "PsyD", "PhD"}:
            corrected_parts.append(part.upper())
        # Capitalize names and titles like "Dr.", "Ms.", "Mr."
        elif part.lower() in {"dr.", "ms.", "mr.", "mrs."}:
            corrected_parts.append(part.capitalize())
        else:
            corrected_parts.append(part.capitalize())

    return " ".join(corrected_parts)

class ReportApp:
    def __init__(self, root):
        root.title("Dr. John P. Shallcross - Psychological Report Generator")
        root.geometry("850x600")

        container = ttk.Frame(root)
        container.pack(fill="both", expand=True)

        canvas = tk.Canvas(container)
        scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        self.root = root
        self.frame = scrollable_frame
        self.entries = {}
        self.transcripts = {"interview": "", "consult": ""}

        # === Audio recording state ===
        self.is_recording = False
        self.is_paused = False  # optional, only needed if using pause/resume later
        self.frames = []
        self.q = queue.Queue()

        self._add_label_entry("Patient Name", "patient")

        ttk.Label(scrollable_frame, text="Title").pack(anchor="w", padx=10)
        self.title_var = tk.StringVar(value="Mr.")
        ttk.Combobox(scrollable_frame, textvariable=self.title_var,
                     values=["Mr.", "Ms.", "Mx.", "Dr.", "The client"], state="readonly").pack(anchor="w", padx=10, pady=2)

        self._add_label_entry("Date of Evaluation (MM/DD/YYYY)", "date")
        self._add_label_entry("Ordered By", "ordered_by")

        ttk.Label(scrollable_frame, text="Tests Administered:").pack(anchor="w", pady=(10, 0))
        self.test_vars = []
        for i, test in enumerate(ALL_TESTS):
            var = tk.BooleanVar(value=(i in DEFAULT_SELECTED))
            cb = tk.Checkbutton(scrollable_frame, text=test.strip("*"), variable=var)
            cb.pack(anchor="w", padx=20)
            self.test_vars.append(var)

        record_controls = ttk.LabelFrame(scrollable_frame, text="Dictation Recording")
        record_controls.pack(padx=10, pady=10, fill="x")

        button_frame = ttk.Frame(record_controls)
        button_frame.pack(pady=5)

        # Start Button
        self.start_btn = ttk.Button(button_frame, text="Start Recording", command=self.start_recording)
        self.start_btn.pack(side="left", padx=5)

        # Stop Button
        self.stop_btn = ttk.Button(button_frame, text="🛑Stop Recording", command=self.stop_recording)
        self.stop_btn.pack(side="left", padx=5)

        consult_frame = ttk.Frame(scrollable_frame)
        consult_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        self.consult_textbox = tk.Text(consult_frame, height=10, wrap="word")
        self.consult_textbox.pack(side="left", fill="both", expand=True)
        consult_scroll = ttk.Scrollbar(consult_frame, orient="vertical", command=self.consult_textbox.yview)
        consult_scroll.pack(side="right", fill="y")
        self.consult_textbox.config(yscrollcommand=consult_scroll.set)

        ttk.Label(scrollable_frame, text="Intellectual/Cognitive Testing:").pack(anchor="w", padx=10, pady=(10, 0))
        
        self._add_label_entry("Full-Scale IQ Score (e.g., 109)", "iq_score")
        self._add_label_entry("Full-Scale IQ Percentile (e.g., 73)", "iq_percentile")


        ttk.Label(scrollable_frame, text="Significant Difference Between Verbal and Performance IQ?").pack(anchor="w", padx=10)
        self.iq_diff_var = tk.StringVar(value="is not")
        ttk.Combobox(scrollable_frame, textvariable=self.iq_diff_var,
                     values=["is", "is not"], state="readonly", width=15).pack(anchor="w", padx=10, pady=(0, 10))
        self._add_label_entry("Years of Education", "education")
#-------------TMT
        ttk.Label(scrollable_frame, text="Trail-Making Test").pack(anchor="w", padx=10, pady=(10, 0))
        self._add_label_entry("Percentile Score (e.g., 25)", "tmt_percentile")
        self.tmt_bottom10_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(scrollable_frame, text='Mark as "bottom 10th"', variable=self.tmt_bottom10_var).pack(anchor="w", padx=10)
        ttk.Label(scrollable_frame, text="Level of Impairment:").pack(anchor="w", padx=10)
        self.tmt_impairment_var = tk.StringVar(value="no")
        ttk.Combobox(scrollable_frame, textvariable=self.tmt_impairment_var,
                     values=["no", "mild", "moderate", "severe", "very severe"], state="readonly", width=20).pack(anchor="w", padx=10, pady=2)
        self.tmt_relative_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(scrollable_frame, text='Add "relative to IQ."', variable=self.tmt_relative_var).pack(anchor="w", padx=10, pady=(0, 10))

#----------WASI
        ttk.Label(scrollable_frame, text="Comparison of WASI-II / WAIS-R:").pack(anchor="w", padx=10, pady=(10, 0))
        ttk.Label(scrollable_frame, text="Immediate Auditory Attention:").pack(anchor="w", padx=10)
        self.waa_var = tk.StringVar(value="unimpaired")
        ttk.Combobox(scrollable_frame, textvariable=self.waa_var,
                     values=["unimpaired", "mild impairment in", "moderate impairment in", "severe impairment in", "very severe impairment in"],
                     state="readonly", width=30).pack(anchor="w", padx=10, pady=(0, 5))
        self.waa_relative_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(scrollable_frame, text='Add "relative to IQ."', variable=self.waa_relative_var).pack(anchor="w", padx=10, pady=(0, 10))
        ttk.Label(scrollable_frame, text="Working Memory:").pack(anchor="w", padx=10)
        self.wm_var = tk.StringVar(value="unimpaired")
        ttk.Combobox(scrollable_frame, textvariable=self.wm_var,
                     values=["unimpaired", "mild impairment in", "moderate impairment in", "severe impairment in", "very severe impairment in"],
                     state="readonly", width=30).pack(anchor="w", padx=10, pady=(0, 10))

        # Letter-Number Sequencing Section
        ttk.Label(scrollable_frame, text="Letter-Number Sequencing:").pack(anchor="w", padx=10, pady=(10, 0))
        self.lns_var = tk.StringVar(value="none")
        ttk.Combobox(scrollable_frame, textvariable=self.lns_var,
                     values=[
                         "not suggestive of impairment",
                         "mild impairment", "moderate impairment",
                         "severe impairment", "very severe impairment"
                     ], state="readonly", width=35).pack(anchor="w", padx=10, pady=(0, 2))
        self.lns_relative_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(scrollable_frame, text='Add "relative to IQ."', variable=self.lns_relative_var).pack(anchor="w", padx=10, pady=(0, 10))

        # Self-report Inventory Section
        ttk.Label(scrollable_frame, text="Self-Report ADHD Inventory:").pack(anchor="w", padx=10, pady=(10, 0))
        self.sr_supportive_var = tk.StringVar(value="no")
        ttk.Combobox(scrollable_frame, textvariable=self.sr_supportive_var,
                     values=["yes", "no"], state="readonly", width=10).pack(anchor="w", padx=10, pady=2)

        ttk.Label(scrollable_frame, text="Presentation Type (if yes):").pack(anchor="w", padx=10)
        self.sr_presentation_var = tk.StringVar()
        ttk.Combobox(scrollable_frame, textvariable=self.sr_presentation_var,
                     values=[
                         "predominately combined presentation",
                         "predominately inattentive presentation",
                         "predominately inattentive presentation of moderate severity",
                         "predominately inattentive presentation with features of hyperactivity",
                         "predominantly inattentive presentation with features of hyperactivity of mild severity",
                         "predominantly hyperactive presentation with features of inattention"
                     ], state="readonly", width=100).pack(anchor="w", padx=10, pady=(0, 10))
        
        # Personality Testing stub — PAI Section Inputs
        ttk.Label(scrollable_frame, text="Results of Objective Personality Testing:").pack(anchor="w", padx=10, pady=(10, 0))

        # Test Type Selector
        self.pai_type_var = tk.StringVar(value="None")
        ttk.Label(scrollable_frame, text="Test Administered:").pack(anchor="w", padx=10)
        ttk.Combobox(scrollable_frame, textvariable=self.pai_type_var,
                    values=["None", "PAI", "PAI-A"], state="readonly", width=10).pack(anchor="w", padx=10, pady=2)

        # Validity Dropdown
        self.pai_validity_var = tk.StringVar(value="Valid (no issues)")
        ttk.Label(scrollable_frame, text="Validity Concerns:").pack(anchor="w", padx=10)
        ttk.Combobox(scrollable_frame, textvariable=self.pai_validity_var,
                    values=["Valid (no issues)", "Infrequency", "Negative Impression", "Positive Impression"],
                    state="readonly", width=25).pack(anchor="w", padx=10, pady=2)
        # Diagnosable Psychopathology Prompt
        ttk.Label(scrollable_frame, text="Is there diagnosable psychopathology?").pack(anchor="w", padx=10, pady=(10, 0))
        self.psychopathology_var = tk.StringVar(value="no")
        ttk.Combobox(scrollable_frame, textvariable=self.psychopathology_var,
                    values=["yes", "no"], state="readonly", width=10).pack(anchor="w", padx=10, pady=2)

        # Clinical Scales and Subscales
        ttk.Label(scrollable_frame, text="Select clinical scales and subscales (if applicable):").pack(anchor="w", padx=10)

        # Dict for storing selected checkboxes
        self.pai_scales = {}

        

        pai_scale_structure = {
            "Somatic": [],
            "Anxiety": ["Cognitive", "Affective", "Physiological"],
            "Anxiety Related Disorders": ["Obsessive-Compulsive", "Phobias", "Traumatic Stress"],
            "Depression": ["Scores > 95T", "Cognitive", "Affective", "Physiological"],
            "Suicidal Ideation": [],
            "Aggression": ["Aggressive Attitude", "Verbal Aggression", "Physical Aggression"],
            "Paranoia": ["Hypervigilance", "Persecution", "Resentment"],
            "Schizophrenia": ["Schizophrenia", "Psychotic Experiences > 85T", "Social Detachment", "Thought Disorder"],
            "Bipolar": ["Activity Level", "Grandiosity", "Irritability"],
            "Borderline": ["Affective Instability", "Identity", "Negative Relationships", "Self-Harm"],
            "Antisocial": ["Scores > 82T", "Antisocial Behaviors", "Egocentricity", "Stimulus-Seeking"],
            "Nonsupport": ["Nonsupport", "Scores > 88T"],
            "Alcohol/Drug": ["Alcohol/Drug", "Scores > 84T"]
        }

        for section, subscales in pai_scale_structure.items():
            ttk.Label(scrollable_frame, text=section + ":").pack(anchor="w", padx=20)
            self.pai_scales[section] = {}
            if not subscales:
                var = tk.BooleanVar()
                self.pai_scales[section]["_"] = var
                ttk.Checkbutton(scrollable_frame, text=section, variable=var).pack(anchor="w", padx=40)
            else:
                for sub in subscales:
                    var = tk.BooleanVar()
                    self.pai_scales[section][sub] = var
                    ttk.Checkbutton(scrollable_frame, text=sub, variable=var).pack(anchor="w", padx=40)


        ttk.Button(scrollable_frame, text="Generate Report", command=self.generate_report).pack(pady=20)
    
    def get_pronoun_replacements(self, title, patient_name):
        last_name = patient_name.split()[-1] if patient_name else "the client"

        pronouns = {
            "Mr.": ("he", "his", "himself", "Mr."),
            "Ms.": ("she", "her", "herself", "Ms."),
            "Mx.": ("they", "their", "themself", "Mx."),
            "Dr.": ("they", "their", "themself", "Dr."),
            "The client": ("they", "their", "themself", "The client")
        }

        he_she, his_her, himself_herself, mr_ms = pronouns.get(
            title, ("they", "their", "themself", "The client")
        )

        replacements = [
            ("[Mr./Ms. Patient last name]", f"{mr_ms} {last_name}"),
            ("Herself", himself_herself.capitalize()),
            ("herself", himself_herself),
            ("Her", his_her.capitalize()),
            ("her", his_her),
            ("She", he_she.capitalize()),
            ("she", he_she),
        ]
        return replacements, f"{mr_ms} {last_name}", he_she, his_her, himself_herself

    def add_heading(self, doc, text, italic=False):
        p = doc.add_paragraph()
        if not text.strip():
            text = " "
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.italic = italic
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        return p

    def add_paragraph(self, doc, text):
        p = doc.add_paragraph()
        if not text.strip():  # If the text is empty or just whitespace, insert a space to make a run
            text = " "

        run = p.add_run(text)
        font = run.font
        font.name = 'Arial'
        font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p
    
    def add_label_paragraph(self, doc, label, content, italic_label=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if not label.strip():
            label = " "
        run_label = p.add_run(label)
        run_label.font.name = 'Arial'
        run_label.font.size = Pt(12)
        run_label.italic = italic_label
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        if not content.strip():
            content = " "
        run_content = p.add_run(content)
        run_content.font.name = 'Arial'
        run_content.font.size = Pt(12)
        run_content._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        return p
    
    
    
    def _add_label_entry(self, label, key):
        ttk.Label(self.frame, text=label).pack(anchor="w", padx=10)
        entry = ttk.Entry(self.frame, width=50)
        entry.pack(anchor="w", padx=10, pady=2)
        self.entries[key] = entry

        #-------------------DICTATION-RECORDING---------

    def audio_callback(self, indata, frames, time_info, status):
        if self.is_recording and not self.is_paused:
            self.q.put(indata.copy())

    def transcribe_and_clean(self, filename):
        try:
            raw = transcribe_audio(audio_file=filename)
            clean = clean_with_together(raw)
            self.transcripts["consult"] = clean
            self.consult_textbox.delete("1.0", tk.END)
            self.consult_textbox.insert(tk.END, clean)
            messagebox.showinfo("✅ Done", "Recording transcribed and cleaned.")
        except Exception as e:
            messagebox.showerror("❌ Error", f"Transcription failed:\n{e}")

    def start_recording(self):
        if self.is_recording:
            messagebox.showwarning("⚠️ Already Recording", "Recording is already in progress.")
            return

        self.is_recording = True
        self.is_paused = False
        self.frames = []
        self.q.queue.clear()

        def threaded_record():
            try:
                with sd.InputStream(samplerate=16000, channels=1, callback=self.audio_callback):
                    while self.is_recording:
                        try:
                            data = self.q.get(timeout=0.1)
                            self.frames.append(data)
                        except queue.Empty:
                            continue
            except Exception as e:
                messagebox.showerror("❌ Audio Error", f"Failed to open audio stream:\n{e}")
                self.is_recording = False

        self.audio_thread = threading.Thread(target=threaded_record, daemon=True)
        self.audio_thread.start()
        messagebox.showinfo("🎙 Recording Started", "Recording has started.")


    def stop_recording(self):
        if not self.is_recording:
            messagebox.showwarning("⚠️ Not Recording", "Recording hasn't started.")
            return

        self.is_recording = False
        self.audio_thread.join()

        if not self.frames:
            messagebox.showerror("❌ No Audio", "No audio was recorded.")
            return

        audio_data = np.concatenate(self.frames, axis=0)

        filename = "consultation.wav"
        try:
            with wave.open(filename, mode='wb') as wf:
                wf.setnchannels(1)
                wf.setsampwidth(2)  # 16-bit PCM
                wf.setframerate(16000)
                wf.writeframes((audio_data * 32767).astype(np.int16).tobytes())
            self.transcribe_and_clean(filename)
        except Exception as e:
            messagebox.showerror("❌ Save Failed", f"Could not save WAV file:\n{e}")

    def record_and_clean(self, section):
        try:
            record_audio()
            raw = transcribe_audio()
            clean = clean_with_together(raw)
            self.transcripts[section] = clean

            if section == "interview":
                self.clinical_textbox.delete("1.0", tk.END)
                self.clinical_textbox.insert(tk.END, clean)
            elif section == "consult":
                self.consult_textbox.delete("1.0", tk.END)
                self.consult_textbox.insert(tk.END, clean)

            messagebox.showinfo("✅ Done", f"{section.title()} transcription completed and cleaned.")
        except Exception as e:
            messagebox.showerror("❌ Error", str(e))

    def get_pronouns(self, title):
        if title == "Mr.":
            return {"he_she": "he", "his_her": "his", "Mr_Ms": "Mr."}
        elif title == "Ms.":
            return {"he_she": "she", "his_her": "her", "Mr_Ms": "Ms."}
        else:
            return {"he_she": "they", "his_her": "their", "Mr_Ms": "The client"}


    def yesno_text(self, var, yes_suffix="", no_suffix=""):
        return yes_suffix if var.get() else no_suffix 
    
    def safe_replace(self, text, old, new):
        # If old contains brackets, use simple string replace
        if "[" in old or "]" in old:
            return text.replace(old, new)
        else:
            # Replace whole words only, case-sensitive
            pattern = r'\b' + re.escape(old) + r'\b'
            return re.sub(pattern, new, text)

    def generate_pai_section(self):
        paragraphs = []

        pai_type = self.pai_type_var.get()
        validity = self.pai_validity_var.get()
        psychopathology = self.psychopathology_var.get()
        title = self.title_var.get()
        patient_name = self.entries["patient"].get()

        pronoun_replacements, full_name, he_she, his_her, himself_herself = self.get_pronoun_replacements(title, patient_name)

        if pai_type not in ["PAI", "PAI-A"]:
            return []

        paragraphs.append(f"{pai_type}.")

        if validity == "Valid (no issues)":
            paragraphs.append(
                f"Results of the PAI are valid. {full_name} responded consistently and there is no indication that "
                f"{he_she} attempted to portray {himself_herself} in a more positive or more negative manner than may actually be the case."
            )
        elif validity == "Infrequency":
            paragraphs.append(f"{full_name} did not attend appropriately to item content in responding to the PAI items.")
        elif validity == "Negative Impression":
            paragraphs.append(
                f"{full_name} presented an extremely negative evaluation of {his_her}self and {his_her} life. "
                f"{he_she.capitalize()} also may be making a “cry for help.” Some deliberate distortion of the clinical picture may also be present."
            )
        elif validity == "Positive Impression":
            paragraphs.append(
                f"{full_name} attempted to portray {his_her}self as exceptionally free of the common shortcomings to which most individuals will admit."
            )
        else:
            paragraphs.append("Validity status not specified.")

        if psychopathology == "no":
            paragraphs.append(
                "All of the clinical scales (Full-Scale Profile, Subscale Profile, and Supplemental Indices) are subclinical. "
                "As such, there is no diagnosable psychopathology."
            )
            return paragraphs

        # Collect and combine subscale paragraphs per domain
        for domain, subs in self.pai_scales.items():
            combined_texts = []
            for label, var in subs.items():
                if var.get():
                    subkey = label if label != "_" else "_"
                    if domain in PAI_PARAGRAPHS and subkey in PAI_PARAGRAPHS[domain]:
                        para = PAI_PARAGRAPHS[domain][subkey]
                        for old, new in pronoun_replacements:
                            para = self.safe_replace(para, old, new)
                        combined_texts.append(para.strip())
                    else:
                        print(f"Missing paragraph for domain '{domain}', subkey '{subkey}'")
            if combined_texts:
                # Join all checked subscale paragraphs for this domain into one paragraph with spaces between
                paragraphs.append(" ".join(combined_texts))

        return paragraphs
    
    def add_paragraphs_from_text(self, doc, text):
        if not text:
            return
        for para in text.strip().split("\n\n"):
            if para.strip():
                self.add_paragraph(doc, para.strip())

    def generate_tmt_section(self):
        raw_value = self.entries.get("tmt_percentile", tk.StringVar()).get().strip()
        tmt_percentile = ordinal(raw_value) if raw_value.isdigit() else raw_value

        tmt_impairment = self.tmt_impairment_var.get()
        bottom10 = self.tmt_bottom10_var.get()
        relative = self.yesno_text(self.tmt_relative_var, " relative to IQ.", ".")

        last_name = self.entries["patient"].get().split()[-1]
        title = self.title_var.get()

        if bottom10:
            return (
                f"In the Trail Making Test (a test of speed for visual search, attention, and mental flexibility), "
                f"{title} {last_name} scored in the bottom 10th percentile range on the more complex task, "
                f"showing {tmt_impairment} impairment{relative}"
            )
        elif raw_value and tmt_impairment:
            return (
                f"In the Trail Making Test (a test of speed for visual search, attention, and mental flexibility), "
                f"{title} {last_name} scored at the {tmt_percentile} percentile range on the more complex task, "
                f"showing {tmt_impairment} impairment{relative}"
            )

        return "Trail Making Test data was not provided."

    def generate_report(self):
        iq_percentile_raw = self.entries["iq_percentile"].get().strip()
      

        # Sanitize into ordinal if valid
        iq_percentile = ordinal(iq_percentile_raw) if iq_percentile_raw.isdigit() else iq_percentile_raw

        patient = self.entries["patient"].get().strip()
        date = self.entries["date"].get().strip()
        ordered_by = clean_ordered_by(self.entries["ordered_by"].get())
        consult = self.consult_textbox.get("1.0", tk.END).strip()

        if not (patient and date and ordered_by and consult):
            messagebox.showwarning("⚠️ Missing Info", "Please fill all required fields and complete both dictations.")
            return

        tests = [t.lstrip("*") for t, var in zip(ALL_TESTS, self.test_vars) if var.get()]
        iq_score = self.entries["iq_score"].get().strip()
        iq_diff = self.iq_diff_var.get()
        education = self.entries["education"].get().strip()
        title = self.title_var.get()
        last_name = patient.split()[-1] if patient else "the client"

        replacements, mr_ms, he_she, his_her, himself_herself = self.get_pronoun_replacements(title, patient)

        
    

        # Section generators
        def generate_iq_section():
            try:
                iq_val = int(iq_score)
            except ValueError:
                return "IQ score must be a number."

            range_text = (
                "Extremely Low Range" if iq_val <= 69 else
                "Borderline Range" if iq_val <= 79 else
                "Low Average Range" if iq_val <= 89 else
                "Average Range" if iq_val <= 109 else
                "High Average Range" if iq_val <= 119 else
                "Superior Range" if iq_val <= 129 else
                "Very Superior Range"
            )

            iq_text = f"{iq_val} ({range_text}"
            if iq_percentile:
                iq_text += f", {iq_percentile} percentile"
            iq_text += ")"

            return (
                f"Results of the Wechsler Abbreviated Scale of Intelligence-II reveal a Full-Scale IQ of {iq_text}. "
                f"There {iq_diff} a significant difference between Verbal and Performance IQ scores. "
                f"{title} {last_name} has {education} years of education."
            )


        def generate_wasi_section():
            waa = self.waa_var.get()
            waa_relative = self.yesno_text(self.waa_relative_var, " relative to IQ")
            wm = self.wm_var.get()
            if waa and wm:
                return (
                    f"Results of a comparison between Full-Scale IQ as measured by the WASI-II and Digit Span of the WAIS-R "
                    f"are suggestive of {waa} immediate auditory attention{waa_relative} and {wm} working memory."
                )
            return "WASI-II / WAIS-R data not provided."

        def generate_lns_section():
            lns_val = self.lns_var.get()
            lns_relative = self.yesno_text(self.lns_relative_var, " relative to IQ.", ".")
            if lns_val == "not suggestive of impairment":
                return "Results of Letter-Number Sequencing are not suggestive of impairment in working memory."
            return f"Results of Letter-Number Sequencing suggest {lns_val} in working memory{lns_relative}"

        def generate_sr_section():
            if self.sr_supportive_var.get() == "no":
                return "Results of a self-report ADHD inventory are not supportive of a diagnosis of ADHD."
            pres = self.sr_presentation_var.get()
            if pres:
                return f"Results of a self-report ADHD inventory are supportive of a diagnosis of ADHD, {pres}."
            return ""
        
        def generate_diagnostic_impression():
            # Hardcoded ICD-10 or DSM code and base diagnosis
            base_diagnosis = "314.00 – Attention-deficit/hyperactivity disorder"

            # Get dropdown input string
            subtype = self.sr_presentation_var.get()  # e.g. "predominantly inattentive presentation with features of hyperactivity"

            # Compose full diagnostic impression
            if subtype:
                diagnostic_impression = f"{base_diagnosis}, {subtype}"
            else:
                diagnostic_impression = base_diagnosis

            # Return or print as paragraph in doc
            return diagnostic_impression


        def generate_summary():
            parts = []
            if iq_score:
                try:
                    iq_val = int(iq_score)
                    range_text = (
                        "Extremely Low Range" if iq_val <= 69 else
                        "Borderline Range" if iq_val <= 79 else
                        "Low Average Range" if iq_val <= 89 else
                        "Average Range" if iq_val <= 109 else
                        "High Average Range" if iq_val <= 119 else
                        "Superior Range" if iq_val <= 129 else
                        "Very Superior Range"
                    )
                    parts.append(
                        f"Results of the Wechsler Abbreviated Scale of Intelligence-II reveal a Full-Scale IQ of {iq_val} ({range_text})."
                    )
                except ValueError:
                    parts.append("IQ score was not a valid number.")



            if iq_diff:
                parts.append(f'Findings did {"not " if iq_diff == "is not" else ""}indicate a significant differential between Verbal and Performance IQ scores.')

            if self.tmt_impairment_var.get() != "no":
                sentence = f"{title} {last_name} showed {self.tmt_impairment_var.get()} impairment in speed for visual search, attention, and mental flexibility"
                if self.tmt_relative_var.get():
                    sentence += " relative to IQ"
                parts.append(sentence + ".")
            parts.append(generate_sr_section())
            if self.lns_var.get() != "not suggestive of impairment":
                lns = f"Letter-Number Sequencing indicates {self.lns_var.get()} in working memory"
                if self.lns_relative_var.get():
                    lns += " relative to IQ"
                parts.append(lns + ".")
            if self.waa_var.get():
                wm = self.wm_var.get()
                parts.append(
                    f"Results of a comparison between Full-Scale IQ and a measure of auditory attention are suggestive of {self.waa_var.get()} immediate auditory attention"
                    + (" relative to IQ" if self.waa_relative_var.get() else "") + f" and {wm} working memory."
                )
            return " ".join(parts)

        # Generate Word doc
        filename = f"{patient.replace(' ', '_').lower()}_report.docx"
        contact_lines = [
            "John P. Shallcross, Psy.D., P.C.",
            "Phone:  770-663-0923",
            "FAX:  770-663-6256",
            "jpshallcross1@gmail.com",
            "drjohnshallcross.com"
        ]
        
        
        try:
            doc = Document()

# Add plain text contact info, centered
      
            for line in contact_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run(line)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self.add_paragraph(doc, "")
            self.add_paragraph(doc, "") 
            self.add_heading(doc, "PSYCHOLOGICAL CONSULT", italic=False).alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, f"PATIENT: {patient}")
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, f"Date of Evaluation: {date}")
            self.add_paragraph(doc, f"Ordered by: {ordered_by}")
            self.add_paragraph(doc, "Reason for Consult: Assist in evaluation of ADHD and differential diagnosis.")
            self.add_paragraph(doc, "")
            self.add_heading(doc, "Tests Administered:", italic=False)
            p = doc.add_paragraph()
            for i, t in enumerate(tests):
                run = p.add_run(t)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                if i != len(tests) - 1:
                    run.add_break()  # inserts a line break instead of new paragraph
            self.add_heading(doc, "Clinical Interview:", italic=False)
            self.add_paragraph(doc, "")
            self.add_heading(doc, "Consultation Findings:", italic=False)
            self.add_paragraph(doc, consult)
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, "")
            self.add_heading(doc, "Results of Intellectual/Cognitive Testing:", italic=True)
            self.add_paragraph(doc, generate_iq_section())
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, self.generate_tmt_section())
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, generate_wasi_section())
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, generate_lns_section())
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, generate_sr_section())
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, "")
            self.add_heading(doc, "Results of Objective Personality Testing:", italic=True)

            for para in self.generate_pai_section():
                if para.strip():
                    self.add_paragraph(doc, para.strip())
                    self.add_paragraph(doc, "")  # space between paragraphs
            self.add_paragraph(doc, "")
            self.add_heading(doc, "Summary/Recommendations", italic=False)
            self.add_paragraph(doc, generate_summary())
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, "")
            self.add_heading(doc, "Diagnostic Impression:", italic=False)
            diagnostic_text = generate_diagnostic_impression()
            self.add_paragraph(doc, diagnostic_text)
            self.add_paragraph(doc, "")
            self.add_heading(doc, "Results of objective personality testing are valid and reveal …", italic=False)
            self.add_paragraph(doc, "")
            self.add_heading(doc, "Signature:", italic=False)
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, "")
            self.add_paragraph(doc, "")
            self.add_heading(doc, "John P. Shallcross, Psy.D.\nLicensed Clinical Psychologist", italic=False)
            superscript_ordinals_in_doc(doc)
            doc.save(filename)
            messagebox.showinfo("✅ Report Generated", f"Report saved as:\n{filename}")
        except Exception as e:
            print("Error saving document:", e)

            messagebox.showerror("❌ Error", f"Could not save document:\n{e}")
if __name__ == "__main__":
    root = tk.Tk()
    app = ReportApp(root)
    root.mainloop()

